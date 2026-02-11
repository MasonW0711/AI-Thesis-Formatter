from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.adapters.ai_classifier import ParagraphAIClassifier
from app.models.schemas import REQUIRED_FONT_NAME, ParagraphRule, RuleSet


CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百千\d]+章")
SECTION_RE = re.compile(r"^第[一二三四五六七八九十百千\d]+節")
SUBSECTION_RE = re.compile(r"^\d+(?:\.\d+)+")
FIGURE_RE = re.compile(r"^(圖|Figure)\s*\d+", re.IGNORECASE)
TABLE_RE = re.compile(r"^(表|Table)\s*\d+", re.IGNORECASE)
TOC_RE = re.compile(r"^(目錄|table of contents|圖目錄|表目錄)$", re.IGNORECASE)
ABSTRACT_RE = re.compile(r"^(摘要|abstract)$", re.IGNORECASE)
KEYWORDS_RE = re.compile(r"^(關鍵詞|keywords)\s*[:：]", re.IGNORECASE)


class FormatApplier:
    """套用頁面與段落格式，並可選擇使用 AI 強化段落語義分類。"""

    def apply(
        self,
        source_docx: Path,
        output_docx: Path,
        ruleset: RuleSet,
        ai_options: dict[str, Any] | None = None,
    ) -> list[str]:
        document = Document(str(source_docx))
        warnings: list[str] = []

        self._apply_page_rule(document, ruleset)
        self._apply_paragraph_rules(document, ruleset, warnings=warnings, ai_options=ai_options)
        self._ensure_index_pages(document, ruleset)

        document.save(str(output_docx))
        return warnings

    def _apply_page_rule(self, document: Document, ruleset: RuleSet) -> None:
        page = ruleset.page

        for idx, section in enumerate(document.sections):
            section.page_width = Pt(page.page_width_pt)
            section.page_height = Pt(page.page_height_pt)
            section.top_margin = Pt(page.margin_top_pt)
            section.bottom_margin = Pt(page.margin_bottom_pt)
            section.left_margin = Pt(page.margin_left_pt)
            section.right_margin = Pt(page.margin_right_pt)
            section.header_distance = Pt(page.header_distance_pt)
            section.footer_distance = Pt(page.footer_distance_pt)
            section.gutter = Pt(page.gutter_pt)

            sect_pr = section._sectPr
            pg_num_nodes = sect_pr.xpath("./w:pgNumType")
            pg_num = pg_num_nodes[0] if pg_num_nodes else None
            if pg_num is None:
                pg_num = OxmlElement("w:pgNumType")
                sect_pr.append(pg_num)

            if page.page_number_format and page.page_number_format != "none":
                pg_num.set(qn("w:fmt"), page.page_number_format)
            if idx == 0 and page.page_number_start > 0:
                pg_num.set(qn("w:start"), str(page.page_number_start))

            self._ensure_footer_page_field(section)

    def _apply_paragraph_rules(
        self,
        document: Document,
        ruleset: RuleSet,
        warnings: list[str],
        ai_options: dict[str, Any] | None = None,
    ) -> None:
        groups = ruleset.groups
        rows: list[dict[str, Any]] = []
        ai_inputs: list[dict[str, Any]] = []

        for idx, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            heuristic_group = self._classify(idx, text, paragraph)
            locked_group = self._locked_group(text)
            p_pr = paragraph._p.pPr
            is_numbered = bool(p_pr is not None and p_pr.numPr is not None)

            row = {
                "index": idx,
                "paragraph": paragraph,
                "heuristic_group": heuristic_group,
                "locked_group": locked_group,
            }
            rows.append(row)

            if not locked_group:
                ai_inputs.append(
                    {
                        "index": idx,
                        "text": text,
                        "heuristic": heuristic_group,
                        "style_name": (paragraph.style.name if paragraph.style else "") or "",
                        "alignment": self._alignment_key(paragraph.alignment),
                        "is_numbered": is_numbered,
                    }
                )

        ai_labels: dict[int, str] = {}
        if ai_inputs:
            classifier = ParagraphAIClassifier.from_overrides(ai_options)
            ai_labels, ai_notes = classifier.classify(ai_inputs)
            warnings.extend(ai_notes)

        for row in rows:
            group = row["locked_group"] or ai_labels.get(row["index"], row["heuristic_group"])
            rule = groups.get(group, groups["body"])
            self._apply_rule_to_paragraph(row["paragraph"], rule)

    def _ensure_index_pages(self, document: Document, ruleset: RuleSet) -> None:
        if not document.paragraphs:
            return

        existing_titles = {p.text.strip() for p in document.paragraphs if p.text and p.text.strip()}

        blocks: list[tuple[str, str]] = []
        if "目錄" not in existing_titles:
            blocks.append(("目錄", 'TOC \\o "1-3" \\h \\z \\u'))
        if "圖目錄" not in existing_titles:
            blocks.append(("圖目錄", 'TOC \\h \\z \\c "Figure"'))
        if "表目錄" not in existing_titles:
            blocks.append(("表目錄", 'TOC \\h \\z \\c "Table"'))

        if not blocks:
            return

        anchor = self._find_first_chapter_paragraph(document) or document.paragraphs[0]
        toc_rule = ruleset.groups.get("toc", ruleset.groups["section_title"])

        # insert in reverse order because insert_paragraph_before always inserts right before anchor
        for title, field_instruction in reversed(blocks):
            page_break_para = anchor.insert_paragraph_before()
            page_break_para.add_run().add_break(WD_BREAK.PAGE)

            field_para = anchor.insert_paragraph_before()
            self._add_complex_field(field_para, field_instruction)
            self._apply_rule_to_paragraph(field_para, ruleset.groups.get("body", toc_rule))

            heading_para = anchor.insert_paragraph_before(title)
            self._apply_rule_to_paragraph(heading_para, toc_rule)

        # Remove possible extra leading page break at very beginning.
        first = document.paragraphs[0]
        if not first.text.strip() and len(first.runs) == 1:
            run_xml = first.runs[0]._element.xml
            if "w:br" in run_xml:
                first._element.getparent().remove(first._element)

    def _apply_rule_to_paragraph(self, paragraph, rule: ParagraphRule) -> None:
        paragraph.alignment = self._to_alignment(rule.alignment)
        paragraph.paragraph_format.line_spacing = rule.line_spacing
        paragraph.paragraph_format.space_before = Pt(rule.space_before_pt)
        paragraph.paragraph_format.space_after = Pt(rule.space_after_pt)
        paragraph.paragraph_format.first_line_indent = Pt(rule.first_line_indent_pt)

        if not paragraph.runs:
            paragraph.add_run()

        enforced_font_name = REQUIRED_FONT_NAME
        for run in paragraph.runs:
            run.font.name = enforced_font_name
            run.font.size = Pt(rule.font_size_pt)
            run.font.bold = rule.bold
            run.font.italic = rule.italic

            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.get_or_add_rFonts()
            r_fonts.set(qn("w:eastAsia"), enforced_font_name)
            r_fonts.set(qn("w:ascii"), enforced_font_name)
            r_fonts.set(qn("w:hAnsi"), enforced_font_name)

    def _classify(self, index: int, text: str, paragraph) -> str:
        locked_group = self._locked_group(text)
        if locked_group:
            return locked_group
        if CHAPTER_RE.match(text):
            return "chapter_title"
        if SECTION_RE.match(text):
            return "section_title"
        if SUBSECTION_RE.match(text):
            return "subsection_title"
        if ABSTRACT_RE.match(text) or KEYWORDS_RE.match(text):
            return "front_matter"

        if index < 25 and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "cover"
        if index < 35 and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "front_matter"

        p_pr = paragraph._p.pPr
        if p_pr is not None and p_pr.numPr is not None:
            return "subsection_title"

        return "body"

    @staticmethod
    def _locked_group(text: str) -> str | None:
        if TOC_RE.match(text):
            return "toc"
        if FIGURE_RE.match(text):
            return "figure_caption"
        if TABLE_RE.match(text):
            return "table_caption"
        return None

    @staticmethod
    def _alignment_key(alignment) -> str:
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "center"
        if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return "right"
        if alignment in {WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.DISTRIBUTE}:
            return "justify"
        return "left"

    @staticmethod
    def _to_alignment(alignment: str):
        if alignment == "center":
            return WD_ALIGN_PARAGRAPH.CENTER
        if alignment == "right":
            return WD_ALIGN_PARAGRAPH.RIGHT
        if alignment == "justify":
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        return WD_ALIGN_PARAGRAPH.LEFT

    @staticmethod
    def _find_first_chapter_paragraph(document: Document):
        for paragraph in document.paragraphs:
            if CHAPTER_RE.match(paragraph.text.strip()):
                return paragraph
        return None

    @staticmethod
    def _ensure_footer_page_field(section) -> None:
        footer = section.footer
        paragraphs = footer.paragraphs
        if not paragraphs:
            paragraph = footer.add_paragraph()
        else:
            paragraph = paragraphs[0]

        has_page_field = False
        for instr in paragraph._element.findall(".//w:instrText", paragraph._element.nsmap):
            if instr.text and "PAGE" in instr.text.upper():
                has_page_field = True
                break

        if not has_page_field:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            FormatApplier._add_complex_field(paragraph, "PAGE \\* MERGEFORMAT")

    @staticmethod
    def _add_complex_field(paragraph, instruction: str) -> None:
        run_begin = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run_begin._r.append(fld_char_begin)

        run_instr = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instr.text = instruction
        run_instr._r.append(instr)

        run_separate = paragraph.add_run()
        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(qn("w:fldCharType"), "separate")
        run_separate._r.append(fld_char_sep)

        paragraph.add_run("1")
        run_end = paragraph.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_char_end)
