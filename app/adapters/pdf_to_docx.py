from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_BREAK


@dataclass
class PdfConversionResult:
    output_docx: Path
    confidence: float
    warning_message: str | None = None


class PdfToDocxAdapter:
    """Convert PDF text content into an intermediate DOCX document."""

    def convert(self, pdf_path: Path, output_docx: Path) -> PdfConversionResult:
        document = Document()

        page_count = 0
        pages_with_text = 0
        extracted_chars = 0
        warning_message: str | None = None

        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages:
                page_count += 1
                text = page.extract_text() or ""
                lines = [line.strip() for line in text.splitlines() if line.strip()]

                if lines:
                    pages_with_text += 1

                if not lines:
                    # Keep page structure even when extraction is empty.
                    placeholder = document.add_paragraph("[此頁未能可靠擷取文字內容]")
                    placeholder.runs[0].italic = True
                else:
                    for line in lines:
                        extracted_chars += len(line)
                        document.add_paragraph(line)

                if page_count < len(pdf.pages):
                    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        if page_count == 0:
            confidence = 0.0
            warning_message = "PDF 沒有可讀頁面。"
        else:
            page_ratio = pages_with_text / page_count
            density = min(1.0, extracted_chars / (page_count * 1000))
            confidence = round((page_ratio * 0.7 + density * 0.3), 3)
            if confidence < 0.55:
                warning_message = "PDF 結構較複雜，轉換品質偏低，請在套用前確認規則。"

        document.save(str(output_docx))
        return PdfConversionResult(output_docx=output_docx, confidence=confidence, warning_message=warning_message)
