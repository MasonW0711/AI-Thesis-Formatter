import time
from pathlib import Path

from docx import Document


def _make_target_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("第一章 測試章節")
    doc.add_paragraph("第一節 測試小節")
    doc.add_paragraph("這是 API smoke test 的內文段落。")
    doc.save(path)


def test_api_end_to_end(client, tmp_path: Path):
    reset = client.post("/api/templates/default/reset")
    assert reset.status_code == 200
    template_id = reset.json()["id"]

    target = tmp_path / "target.docx"
    _make_target_docx(target)

    with target.open("rb") as fh:
        create_resp = client.post(
            "/api/jobs",
            data={"template_id": template_id},
            files={
                "target_file": (
                    "target.docx",
                    fh,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert create_resp.status_code == 200
    job_id = create_resp.json()["job_id"]

    final_payload = None
    for _ in range(40):
        status_resp = client.get(f"/api/jobs/{job_id}")
        assert status_resp.status_code == 200
        payload = status_resp.json()
        final_payload = payload
        if payload["status"] in {"success", "failed"}:
            break
        time.sleep(0.1)

    assert final_payload is not None
    assert final_payload["status"] == "success"
    assert final_payload["download_url"]

    download_resp = client.get(final_payload["download_url"])
    assert download_resp.status_code == 200
    assert len(download_resp.content) > 0
