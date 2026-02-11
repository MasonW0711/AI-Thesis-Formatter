from pathlib import Path

from docx import Document

from app.engines.format_applier import FormatApplier
from app.models.schemas import REQUIRED_FONT_NAME, RuleSet


def _build_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("第一章 緒論")
    doc.add_paragraph("第一節 研究背景")
    doc.add_paragraph("1.1.1 研究目的")
    doc.add_paragraph("這是一段測試內文，用來驗證格式套用是否穩定。")
    doc.add_paragraph("圖 1-1 系統架構圖")
    doc.add_paragraph("表 1-1 測試資料表")
    doc.save(path)


def _extract_snapshot(path: Path) -> tuple[list[tuple[str, str | None, float | None, bool | None]], int]:
    doc = Document(path)
    rows = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        run = p.runs[0] if p.runs else None
        font_name = run.font.name if run else None
        size = run.font.size.pt if run and run.font.size else None
        bold = run.font.bold if run else None
        rows.append((p.text.strip(), font_name, size, bold))

    toc_count = sum(1 for p in doc.paragraphs if p.text.strip() == "目錄")
    return rows, toc_count


def test_format_apply_is_idempotent(tmp_path: Path):
    source = tmp_path / "source.docx"
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"

    _build_sample_docx(source)
    rules = RuleSet()

    applier = FormatApplier()
    applier.apply(source, first, rules, ai_options={"provider": "off"})
    applier.apply(first, second, rules, ai_options={"provider": "off"})

    first_snapshot, first_toc_count = _extract_snapshot(first)
    second_snapshot, second_toc_count = _extract_snapshot(second)

    assert first_snapshot == second_snapshot
    assert first_toc_count == 1
    assert second_toc_count == 1
    assert all(row[1] == REQUIRED_FONT_NAME for row in second_snapshot)
