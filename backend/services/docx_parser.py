"""
論文格式調整系統 - Word 文件解析服務
"""
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
import os


@dataclass
class DocxTextBlock:
    """Word 文字區塊"""
    text: str
    font_name: str = ""
    font_size: float = 12.0
    is_bold: bool = False
    is_italic: bool = False
    alignment: str = "left"
    style_name: str = ""


@dataclass
class DocxSection:
    """Word 章節"""
    section_number: int
    margin_top: float = 72.0
    margin_bottom: float = 72.0
    margin_left: float = 72.0
    margin_right: float = 72.0
    page_width: float = 595.0
    page_height: float = 842.0


@dataclass
class DocxStructure:
    """Word 文件結構"""
    filename: str
    paragraph_count: int = 0
    sections: List[DocxSection] = field(default_factory=list)
    text_blocks: List[DocxTextBlock] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocxParser:
    """Word 文件解析器"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc: Optional[DocxDocument] = None
    
    def __enter__(self):
        self.doc = DocxDocument(self.file_path)
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.doc = None
    
    def get_paragraph_count(self) -> int:
        """取得段落數量"""
        if self.doc:
            return len(self.doc.paragraphs)
        return 0
    
    def parse(self) -> Optional[DocxStructure]:
        """解析 Word 文件"""
        if not self.doc:
            return None
        
        structure = DocxStructure(
            filename=os.path.basename(self.file_path),
            paragraph_count=len(self.doc.paragraphs)
        )
        
        # 解析章節資訊
        for i, section in enumerate(self.doc.sections):
            doc_section = DocxSection(
                section_number=i + 1,
                margin_top=self._twips_to_pt(section.top_margin),
                margin_bottom=self._twips_to_pt(section.bottom_margin),
                margin_left=self._twips_to_pt(section.left_margin),
                margin_right=self._twips_to_pt(section.right_margin),
                page_width=self._twips_to_pt(section.page_width),
                page_height=self._twips_to_pt(section.page_height)
            )
            structure.sections.append(doc_section)
        
        # 解析段落
        for para in self.doc.paragraphs:
            if not para.text.strip():
                continue
            
            # 取得段落格式
            font_name = ""
            font_size = 12.0
            is_bold = False
            is_italic = False
            
            if para.runs:
                run = para.runs[0]
                if run.font.name:
                    font_name = run.font.name
                if run.font.size:
                    font_size = run.font.size.pt
                is_bold = run.font.bold or False
                is_italic = run.font.italic or False
            
            # 取得對齊方式
            alignment = "left"
            if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                alignment = "center"
            elif para.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                alignment = "right"
            elif para.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                alignment = "justify"
            
            text_block = DocxTextBlock(
                text=para.text,
                font_name=font_name,
                font_size=font_size,
                is_bold=is_bold,
                is_italic=is_italic,
                alignment=alignment,
                style_name=para.style.name if para.style else ""
            )
            structure.text_blocks.append(text_block)
        
        # 取得文件屬性
        core_props = self.doc.core_properties
        structure.metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else ""
        }
        
        return structure
    
    def _twips_to_pt(self, twips) -> float:
        """將 twips 轉換為 pt"""
        if twips is None:
            return 72.0
        # 1 inch = 1440 twips = 72 pt
        return twips / 20.0
