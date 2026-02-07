"""
論文格式調整系統 - Word 文件生成服務
"""
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Dict, Any, Optional

from services.format_analyzer import LearnedFormatRules, ElementType


def clamp_number(value: Any, low: float, high: float, default: float) -> float:
    """數值夾限，避免極端值導致文件生成失敗"""
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return default
    return max(low, min(high, numeric))


def set_chinese_font(run, font_name: str):
    """安全設定中文字型"""
    run.font.name = font_name
    # 為中文設定 eastAsia 字型
    from docx.oxml.ns import qn
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


class DocxGenerator:
    """Word 文件生成器"""
    
    def __init__(self, output_path: str, learned_rules: Optional[LearnedFormatRules] = None):
        self.output_path = output_path
        self.learned_rules = learned_rules
        self.doc = DocxDocument()
        self._setup_document()
    
    def _setup_document(self):
        """設定文件頁面"""
        section = self.doc.sections[0]
        
        # 設定頁面大小 A4
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        
        # 設定邊距
        if self.learned_rules:
            section.top_margin = Pt(clamp_number(self.learned_rules.margin_top, 0.0, 252.0, 72.0))
            section.bottom_margin = Pt(clamp_number(self.learned_rules.margin_bottom, 0.0, 252.0, 72.0))
            section.left_margin = Pt(clamp_number(self.learned_rules.margin_left, 0.0, 252.0, 72.0))
            section.right_margin = Pt(clamp_number(self.learned_rules.margin_right, 0.0, 252.0, 72.0))
        else:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(2.5)
    
    def get_font_name(self) -> str:
        """取得字型名稱"""
        if self.learned_rules and self.learned_rules.main_font_name:
            return self.learned_rules.main_font_name
        return "標楷體"
    
    def generate(self, elements: List[Dict[str, Any]]) -> str:
        """生成 Word 文件"""
        font_name = self.get_font_name()
        
        for element in elements:
            element_type = element.get("type", ElementType.PARAGRAPH)
            text = element.get("text", "").strip()
            
            if not text:
                continue
            
            # 轉換元素類型
            if isinstance(element_type, str):
                try:
                    element_type = ElementType(element_type)
                except ValueError:
                    element_type = ElementType.PARAGRAPH
            
            # 添加段落
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            
            # 設定字型
            set_chinese_font(run, font_name)
            
            # 根據元素類型設定格式
            if element_type == ElementType.TITLE:
                run.font.size = Pt(clamp_number(self.learned_rules.title_font_size if self.learned_rules else 18, 10.0, 36.0, 18.0))
                run.font.bold = True
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
            elif element_type == ElementType.CHAPTER:
                run.font.size = Pt(clamp_number(self.learned_rules.chapter_font_size if self.learned_rules else 16, 10.0, 30.0, 16.0))
                run.font.bold = True
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
            elif element_type == ElementType.SECTION:
                run.font.size = Pt(clamp_number(self.learned_rules.section_font_size if self.learned_rules else 14, 10.0, 24.0, 14.0))
                run.font.bold = True
                
            elif element_type == ElementType.SUBSECTION:
                run.font.size = Pt(12)
                run.font.bold = True
                
            else:  # PARAGRAPH
                run.font.size = Pt(clamp_number(self.learned_rules.paragraph_font_size if self.learned_rules else 12, 8.0, 24.0, 12.0))
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
                # 設定首行縮排
                if self.learned_rules and self.learned_rules.first_line_indent > 0:
                    para.paragraph_format.first_line_indent = Pt(clamp_number(self.learned_rules.first_line_indent, 0.0, 72.0, 24.0))
                else:
                    para.paragraph_format.first_line_indent = Cm(0.85)
                
                # 設定行距
                if self.learned_rules:
                    para.paragraph_format.line_spacing = clamp_number(self.learned_rules.line_spacing, 1.0, 3.0, 1.5)
                else:
                    para.paragraph_format.line_spacing = 1.5
        
        # 儲存文件
        self.doc.save(self.output_path)
        return self.output_path
    
    def apply_format_to_document(self, source_path: str) -> str:
        """將格式套用到現有 Word 文件"""
        source_doc = DocxDocument(source_path)
        
        font_name = self.get_font_name()
        
        # 設定頁面邊距
        for section in source_doc.sections:
            if self.learned_rules:
                section.top_margin = Pt(clamp_number(self.learned_rules.margin_top, 0.0, 252.0, 72.0))
                section.bottom_margin = Pt(clamp_number(self.learned_rules.margin_bottom, 0.0, 252.0, 72.0))
                section.left_margin = Pt(clamp_number(self.learned_rules.margin_left, 0.0, 252.0, 72.0))
                section.right_margin = Pt(clamp_number(self.learned_rules.margin_right, 0.0, 252.0, 72.0))
        
        # 處理每個段落
        for para in source_doc.paragraphs:
            if not para.text.strip():
                continue
            
            # 判斷段落類型
            is_title = self._is_title_paragraph(para)
            is_heading = self._is_heading_paragraph(para)
            
            for run in para.runs:
                # 設定中文字型
                set_chinese_font(run, font_name)
                
                # 設定字型大小
                if is_title:
                    run.font.size = Pt(clamp_number(self.learned_rules.title_font_size if self.learned_rules else 18, 10.0, 36.0, 18.0))
                    run.font.bold = True
                elif is_heading:
                    run.font.size = Pt(clamp_number(self.learned_rules.chapter_font_size if self.learned_rules else 14, 10.0, 30.0, 14.0))
                    run.font.bold = True
                else:
                    run.font.size = Pt(clamp_number(self.learned_rules.paragraph_font_size if self.learned_rules else 12, 8.0, 24.0, 12.0))
            
            # 設定段落格式
            if not is_title and not is_heading:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                if self.learned_rules and self.learned_rules.first_line_indent > 0:
                    para.paragraph_format.first_line_indent = Pt(clamp_number(self.learned_rules.first_line_indent, 0.0, 72.0, 24.0))
                else:
                    para.paragraph_format.first_line_indent = Cm(0.85)
                
                if self.learned_rules:
                    para.paragraph_format.line_spacing = clamp_number(self.learned_rules.line_spacing, 1.0, 3.0, 1.5)
        
        source_doc.save(self.output_path)
        return self.output_path
    
    def _is_title_paragraph(self, para) -> bool:
        """判斷是否為標題段落"""
        if para.style and "Title" in para.style.name:
            return True
        if para.runs:
            run = para.runs[0]
            if run.font.size and run.font.size.pt >= 16:
                if run.font.bold:
                    return True
        return False
    
    def _is_heading_paragraph(self, para) -> bool:
        """判斷是否為標題段落"""
        if para.style and "Heading" in para.style.name:
            return True
        text = para.text.strip()
        # 檢查常見標題模式
        if any(keyword in text for keyword in ["第", "章", "節", "Chapter", "Section"]):
            if para.runs and para.runs[0].font.bold:
                return True
        return False
